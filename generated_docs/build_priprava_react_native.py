from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "/Users/filip.skrget/Documents/VozimVarno/generated_docs/PripravaNaProjekt_ReactNativeExpo.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_cell_text(cell, bold=False, size=9.5):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.bold = bold


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "F2F4F7")
        style_cell_text(hdr[i], bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            style_cell_text(cells[i])
    set_table_width(table, widths)
    document.add_paragraph()
    return table


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = "Calibri"
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_label_paragraph(document, label, text):
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label)
    set_run_font(r, bold=True)
    r = p.add_run(text)
    set_run_font(r)


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_run_font(r)


doc = Document()
section = doc.sections[0]
section.start_type = WD_SECTION_START.NEW_PAGE
section.page_width = Inches(8.5)
section.page_height = Inches(11)
for margin in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
    setattr(section, margin, Inches(1))
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 16, 8),
    ("Heading 2", 13, "2E74B5", 12, 6),
    ("Heading 3", 12, "1F4D78", 8, 4),
]:
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(2)
r = title.add_run("Priprava na projekt")
set_run_font(r, 22, True, "0B2545")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(2)
r = subtitle.add_run("Tehnologije za vseprisotne aplikacije")
set_run_font(r, 12, False, "555555")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(14)
r = meta.add_run("UM FERI • 2025/2026")
set_run_font(r, 10.5, False, "555555")

doc.add_heading("A. Predstavitev projektne ideje", level=1)
add_label_paragraph(doc, "Naslov ideje: ", "DriveSafe – Pametni sopotnik za voznike")
add_label_paragraph(doc, "Vodja ekipe: ", "Lovro Čuš")

doc.add_heading("Besedni opis ideje", level=2)
add_label_paragraph(
    doc,
    "a) Tematika: ",
    "Aplikacija spada v domeno prometne varnosti in mobilne telematike. Namenjena je spremljanju voznikovega vedenja v realnem času z uporabo senzorjev mobilnega telefona. Projekt je zasnovan kot večplatformska mobilna aplikacija v React Native z Expo orodji, pri čemer se za funkcionalnosti z dostopom do senzorjev uporabi Expo Development Build in po potrebi native moduli.",
)
add_label_paragraph(
    doc,
    "b) Ciljni uporabniki: ",
    "Mladi vozniki začetniki, vozniške šole in starši najstnikov, ki želijo nadzor nad varnostjo vožnje svojih otrok.",
)
add_label_paragraph(
    doc,
    "c) Opis: ",
    "Aplikacija med vožnjo sprotno zbira podatke iz kamere, GPS-a, pospeškometra, giroskopa in mikrofona. Na podlagi teh podatkov zaznava znake utrujenosti, agresivno vožnjo in nevarne manevre. Po vsaki vožnji voznik dobi oceno varnosti in podroben pregled poti, incidentov ter priporočil za varnejšo vožnjo.",
)

doc.add_heading("Ključne funkcionalnosti", level=2)
add_table(
    doc,
    ["Zap. št.", "Naziv funkcionalnosti"],
    [
        ["1.", "Zaznava utrujenosti voznika (kamera + ML)"],
        ["2.", "Analiza načina vožnje (pospeškometer + giroskop)"],
        ["3.", "Ocena varnosti vožnje (scoring algoritem)"],
        ["4.", "Sledenje poti in hitrosti (GPS)"],
        ["5.", "Spletna povezava – nevarne ceste in vreme (REST API)"],
        ["6.", "Pregled zgodovine voženj (lokalna baza)"],
        ["7.", "Zaznava hrupa in jokajočega otroka (mikrofon)"],
        ["8.", "Uporabniški vmesnik in nastavitve (React Native + Expo)"],
    ],
    [1100, 8260],
)

doc.add_heading("Podrobnejši opis funkcionalnosti", level=2)
feature_rows = [
    [
        "1.",
        "Tehnična izvedba: Expo Camera oziroma React Native kamera v Expo Development Build za zajem slike v realnem času; analiza obraza z on-device ML modelom (npr. TensorFlow Lite/MediaPipe prek native modula). Logika obdelave teče asinhrono v ločenih servisnih oziroma task funkcijah, da ne blokira uporabniškega vmesnika.\nPodatki/viri: Prednja kamera telefona, video tok ali periodični zajemi, frekvenca mežikanja, čas zaprtih oči, kot nagiba glave.\nOmejitve: Slabša zanesljivost pri slabi svetlobi ali z očali; visoka poraba baterije pri stalnem zajemu; zahteva dovoljenje CAMERA in dodatno testiranje na fizičnih napravah.",
    ],
    [
        "2.",
        "Tehnična izvedba: Expo Sensors za dostop do pospeškometra in giroskopa; obdelava meritev v TypeScriptu z nastavljivimi pragovi G-sile; dogodki se beležijo s časovnim žigom in povežejo s trenutno vožnjo.\nPodatki/viri: Pospeškometer, giroskop in GPS hitrost za kontekst; pragovi so nastavljivi v nastavitvah aplikacije.\nOmejitve: Natančnost je odvisna od kakovosti senzorjev telefona; vibracije ceste lahko sprožijo lažne alarme; telefon mora biti stabilno nameščen v vozilu.",
    ],
    [
        "3.",
        "Tehnična izvedba: Scoring algoritem v TypeScriptu po koncu vožnje agregira podatke iz lokalne baze; uteži za utrujenost, nevarne manevre, hitrost in hrup izračunajo skupno oceno 0–100; rezultat se shrani skupaj s povzetkom vožnje.\nPodatki/viri: Zabeleženi incidenti, trajanje vožnje, skupna razdalja, povprečna in najvišja hitrost.\nOmejitve: Ocena je hevristična in ne temelji na kalibriranem varnostnem modelu; krajše vožnje pod približno 2 minutama dajo manj zanesljiv rezultat.",
    ],
    [
        "4.",
        "Tehnična izvedba: Expo Location za sledenje GPS koordinatam in hitrosti; koordinate se shranjujejo lokalno vsakih nekaj sekund; pot se po vožnji prikaže z react-native-maps ali MapLibre.\nPodatki/viri: GPS koordinate, hitrost, nadmorska višina in časovni žigi; možnost primerjave s podatki o omejitvah hitrosti, kadar so na voljo.\nOmejitve: GPS ni zanesljiv v predorih ali med visokimi zgradbami; pogosto beleženje povečuje porabo baterije in prostora.",
    ],
    [
        "5.",
        "Tehnična izvedba: REST klici z fetch/Axios in TanStack Query za pridobivanje vremenskih ter cestnih podatkov; Repository vzorec odloča med svežimi spletnimi podatki in lokalnim cacheom; Firebase ali Supabase omogoča sinhronizacijo rezultatov s starši.\nPodatki/viri: Vremenski API, podatki o nevarnih cestnih odsekih, oblačna baza za deljenje rezultatov, lokalni cache za offline delovanje.\nOmejitve: Zahteva internetno povezavo za sveže podatke; oblačne storitve lahko povzročijo stroške; API ključi morajo biti shranjeni varno in ne neposredno v kodi.",
    ],
    [
        "6.",
        "Tehnična izvedba: Expo SQLite za lokalno hrambo voženj, incidentov in GPS točk; podatkovni sloj v TypeScriptu z jasno ločenimi repozitoriji; prikaz zgodovine z FlatList/SectionList in filtriranjem po datumu.\nPodatki/viri: Lokalna SQLite baza na telefonu; podatki iz vseh funkcionalnosti, ki se med vožnjo zapisujejo.\nOmejitve: Baza raste z vsako vožnjo, zato je potrebna arhivacija ali brisanje starih voženj; brez oblačne sinhronizacije podatki niso samodejno dostopni z drugih naprav.",
    ],
    [
        "7.",
        "Tehnična izvedba: Expo Audio oziroma native audio modul v Expo Development Build za zajem zvoka; analiza amplitude za zaznavo prekoračitve glasnosti; po potrebi uporaba lahkega ML modela za prepoznavo jokajočega otroka.\nPodatki/viri: Mikrofon telefona, zvočni vzorci, amplitude in klasifikacijski rezultati.\nOmejitve: Zahteva dovoljenje RECORD_AUDIO oziroma Microphone; glasba ali radio v avtu zmanjša zanesljivost zaznave; stalno vzorčenje povečuje porabo procesorja in baterije.",
    ],
    [
        "8.",
        "Tehnična izvedba: Uporabniški vmesnik v React Native komponentah z Expo Router navigacijo; nastavitve se shranjujejo z AsyncStorage/SecureStore, stanje aplikacije pa se vodi z React Context, Zustand ali podobno knjižnico; podpora temni/svetli temi in lokalizaciji.\nPodatki/viri: Nastavitve pragov, jezika in teme; podatki iz lokalne baze za prikaz statistik; prevodi za slovenščino in angleščino.\nOmejitve: Razlike med iOS in Android dovoljenji zahtevajo ločeno testiranje; animacije in stalni senzorji lahko upočasnijo delovanje na šibkejših napravah.",
    ],
]
add_table(doc, ["Zap. št.", "Podroben opis funkcionalnosti"], feature_rows, [800, 8560])

doc.add_heading("Predlagani tehnološki sklad", level=2)
add_bullets(
    doc,
    [
        "React Native + Expo kot osnovni okvir za večplatformski razvoj mobilne aplikacije.",
        "Expo Development Build/EAS Build za funkcionalnosti, ki zahtevajo dodatne native module.",
        "TypeScript za tipno varnejšo poslovno logiko, scoring algoritem in podatkovni sloj.",
        "Expo Sensors, Expo Location, Expo Camera, Expo Audio in Expo SQLite za dostop do ključnih funkcionalnosti naprave.",
        "Firebase ali Supabase za opcijsko sinhronizacijo rezultatov in pregled s strani staršev.",
    ],
)

doc.add_heading("B. Predstavitev ekipe", level=1)
add_table(
    doc,
    ["Član", "Ime in priimek", "IDUM", "Odgovornost za funkcionalnosti"],
    [
        ["Član 1", "Lovro Čuš", "1002519705", "1, 2, 3"],
        ["Član 2", "Filip Škrget", "1002507146", "4, 5, 3"],
        ["Član 3", "Marko Kramer", "1002514681", "6, 7, 8"],
    ],
    [1300, 3000, 2300, 2760],
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run("Pripravljeno kot React Native/Expo različica izvorne projektne priprave.")
set_run_font(r, 9.5, False, "555555")

doc.save(OUT)
print(OUT)
